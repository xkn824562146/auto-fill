"""Word 模板扫描器：识别表格中的空字段及其坐标，保存原始格式。"""

import logging
from dataclasses import dataclass, field

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


@dataclass
class CellFormat:
    """单元格格式信息，用于回填时保持一致。"""
    bold: bool | None = None
    italic: bool | None = None
    underline: bool | None = None
    font_size: object = None
    font_name: str | None = None
    color_rgb: object = None
    alignment: object = None


@dataclass
class FieldSlot:
    """一个待填充的字段位置。"""
    table_idx: int
    row_idx: int
    cell_idx: int
    label: str
    standard: str = ""  # 标准要求文本（如 ≤0.16）
    fmt: CellFormat = field(default_factory=CellFormat)


class WordScanner:
    """扫描 .docx 模板表格，找出所有待填充的空字段。"""

    def scan(self, template_path: str) -> list[FieldSlot]:
        """扫描模板，返回待填充字段列表。"""
        doc = Document(template_path)
        slots: list[FieldSlot] = []

        for table_idx, table in enumerate(doc.tables):
            table_slots = self._scan_table(table, table_idx)
            slots.extend(table_slots)

        return slots

    def _scan_table(self, table, table_idx: int) -> list[FieldSlot]:
        """扫描单个表格，识别空字段。"""
        slots: list[FieldSlot] = []
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        _BLANK_MARKERS = {"以下空白", "以下为空白", "以下为空"}

        for row_idx, row in enumerate(table.rows):
            # 遇到"以下空白"标记行，停止扫描后续行
            row_texts = {cell.text.strip() for cell in row.cells}
            if row_texts & _BLANK_MARKERS:
                break

            for cell_idx, cell in enumerate(row.cells):
                # 跳过非空单元格
                text = cell.text.strip()
                if text:
                    continue

                # 跳过合并单元格中被"吃掉"的位置（与左侧或上方共享同一个 _tc）
                if self._is_merged_placeholder(table, row_idx, cell_idx):
                    continue

                # 在同行或同列找标签
                label = self._find_label(table, row_idx, cell_idx, num_rows, num_cols)
                if not label:
                    continue

                # 提取该行的标准要求文本
                standard = self._extract_standard(row, num_cols)

                # 保存格式
                fmt = self._capture_format(cell)

                slots.append(FieldSlot(
                    table_idx=table_idx,
                    row_idx=row_idx,
                    cell_idx=cell_idx,
                    label=label,
                    standard=standard,
                    fmt=fmt,
                ))

        return slots

    def _extract_standard(self, row, num_cols: int) -> str:
        """从行中提取标准要求文本（含比较运算符或 / 的单元格，去重）。"""
        import re
        seen = set()
        parts = []
        for c in range(num_cols):
            text = row.cells[c].text.strip()
            if text and (re.search(r'[≤≥<>=±＞﹤≧≦﹥]', text) or text == '/') and text not in seen:
                seen.add(text)
                parts.append(text)
        return " ".join(parts) if parts else ""

    def _is_merged_placeholder(self, table, row_idx: int, cell_idx: int) -> bool:
        """检查单元格是否是合并区域中的占位符（非主单元格）。"""
        row = table.rows[row_idx]
        cell = row.cells[cell_idx]
        tc = cell._tc

        # 同一行中，如果左侧单元格共享同一个 _tc，则当前是占位符
        if cell_idx > 0 and row.cells[cell_idx - 1]._tc is tc:
            return True

        # 同一列中，如果上方单元格共享同一个 _tc，则当前是占位符
        if row_idx > 0 and table.rows[row_idx - 1].cells[cell_idx]._tc is tc:
            return True

        return False

    def _find_label(self, table, row_idx: int, cell_idx: int, num_rows: int, num_cols: int) -> str | None:
        """在同行或同列寻找标签文字。

        策略：
        1. 找到该表格的表头行（通过预计算缓存）
        2. 表头行内的空单元格：向左找同行标签
        3. 数据行的空单元格：从表头行向下到当前行上方，找最近的非空单元格作为列标题；
           同时向左找行标签；两者拼接成完整标签
        """
        row = table.rows[row_idx]

        # 获取表头行索引和检验项目列索引
        header_row, item_col_idx = self._get_header_row(table, num_rows, num_cols)
        if header_row is not None and item_col_idx is None:
            logger.warning("[_find_label] 表头行=%d 但未找到'检验项目'/'检测项目'列，回退到向左扫描", header_row)

        if row_idx == header_row:
            # 表头行内：向左找同行标签
            for c in range(cell_idx - 1, -1, -1):
                text = row.cells[c].text.strip()
                if text and len(text) > 1:
                    return text
            return None

        # 数据行：直接从表头行取列标题，再向左找行标签
        col_header = None
        if header_row is not None:
            hrow = table.rows[header_row]
            # 先检查表头行当前列
            hcell = hrow.cells[cell_idx]
            text = hcell.text.strip()
            if text and len(text) > 1:
                # 检查该表头是否是合并单元格的占位符（跨列合并）
                # 如果是占位符，说明当前列有自己的列标题在右边，不应使用合并的列标题
                tc = hcell._tc
                is_placeholder = cell_idx > 0 and hrow.cells[cell_idx - 1]._tc is tc
                if not is_placeholder:
                    col_header = ''.join(text.split()).strip()
            if col_header is None:
                # 向左找最近的非空且非合并占位符的表头单元格
                for c in range(cell_idx - 1, -1, -1):
                    hcell = hrow.cells[c]
                    text = hcell.text.strip()
                    if text and len(text) > 1:
                        # 跳过合并占位符
                        tc = hcell._tc
                        if c > 0 and hrow.cells[c - 1]._tc is tc:
                            continue
                        col_header = ''.join(text.split()).strip()
                        break

        # 优先从检验项目列读取行标签，再扫描子项，找不到再向左扫描
        row_label = None
        row_label_col = None
        if item_col_idx is not None and item_col_idx != cell_idx:
            text = row.cells[item_col_idx].text.strip()
            if text and len(text) > 1:
                row_label = text
                row_label_col = item_col_idx
                # 扫描 item_col 和空单元格之间的列，拼接有意义的文本
                added = {text}
                for c in range(item_col_idx + 1, cell_idx):
                    mid_text = row.cells[c].text.strip().replace('\n', '')
                    if (mid_text
                            and mid_text not in added
                            and not self._is_spec_text(mid_text)
                            and not self._is_unit_label(mid_text)):
                        row_label = f"{row_label}-{mid_text}"
                        added.add(mid_text)
        if row_label is None:
            for c in range(cell_idx - 1, -1, -1):
                text = row.cells[c].text.strip()
                if text and len(text) > 1 and not self._is_spec_text(text) and not self._is_unit_label(text):
                    row_label = text
                    row_label_col = c
                    break

        # 调试：打印找到的标签和整行内容
        row_all = [row.cells[c].text.strip()[:15] for c in range(num_cols)]
        logger.info("[_find_label] row=%d col=%d -> row_label=%r row_label_col=%s | row=%s",
                    row_idx, cell_idx, row_label, row_label_col, row_all)

        # 如果行标签是子项（如"经向"、"纬向"、"横向"、"纵向"等），
        # 在同行向左查找父级标签（如"拉伸断裂强力"），找不到再向上查找
        if row_label and self._is_sub_item(row_label) and row_label_col is not None:
            parent_label = self._find_parent_label(table, row_idx, row_label_col, header_row, num_cols)
            logger.info("[_find_label] sub_item=%r parent=%r", row_label, parent_label)
            if parent_label:
                row_label = f"{parent_label}-{row_label}"

        # 组合标签：必须有行标签（检测项名称），列标题只是辅助
        if col_header and row_label:
            row_is_unit = self._is_unit_label(row_label)
            col_is_unit = self._is_unit_label(col_header)
            # 优先保留方向子项（经向/纬向），丢弃单位子项（N/50mm/%）
            if row_is_unit and not col_is_unit:
                return col_header
            if col_is_unit and not row_is_unit:
                return row_label
            if row_is_unit and col_is_unit:
                return row_label
            return f"{row_label}-{col_header}"
        if row_label:
            return row_label
        return None

    @staticmethod
    def _is_sub_item(text: str) -> bool:
        """判断文本是否为子项标签（方向、位置等）。"""
        sub_items = {
            "经向", "纬向", "横向", "纵向",
            "MD向", "CD向",
            "经向保留率", "纬向保留率",
        }
        return text in sub_items

    @staticmethod
    def _is_unit_label(text: str) -> bool:
        """判断文本是否为单位标签（如 MPa、N/50mm、mm、%、℃ 等）。"""
        import re
        # 含 / 的单位（如 N/50mm、W/(m·K)、MJ/kg、g/m2、g/㎡）
        if re.search(r'[A-Za-z㎡]+/', text) or re.search(r'/[A-Za-z㎡]', text):
            return True
        # 纯单位符号
        if re.match(r'^[°%℃㎡]+$', text):
            return True
        # 常见纯字母单位（大小写均可，长度 ≤ 5）
        if re.match(r'^[A-Za-z]{1,5}$', text):
            return True
        # 含数字的单位变体（如 m2、m3、cm2）
        if re.match(r'^[A-Za-z]+\d$', text) and len(text) <= 5:
            return True
        return False

    def _find_parent_label(self, table, row_idx: int, col_idx: int, header_row: int | None, num_cols: int) -> str | None:
        """查找子项（如"经向"）的父级标签（如"拉伸断裂强力"）。

        策略：
        1. 先在同行向左查找（父级标签通常在子项左侧的列中）
        2. 如果没找到，向上扫描同一列（处理合并行的情况）
        """
        row = table.rows[row_idx]

        # 策略1：在同行向左查找父级标签
        for c in range(col_idx - 1, -1, -1):
            text = row.cells[c].text.strip()
            logger.info("[_find_parent] left row=%d col=%d text=%r is_sub=%s is_spec=%s",
                        row_idx, c, text[:20] if text else '', self._is_sub_item(text), self._is_spec_text(text) if text else False)
            if text and len(text) > 1 and not self._is_spec_text(text):
                if not self._is_sub_item(text):
                    logger.info("[_find_parent] found LEFT parent=%r", text)
                    return text

        # 策略2：向上扫描同一列（处理合并行）
        stop_row = header_row if header_row is not None else -1
        for r in range(row_idx - 1, stop_row, -1):
            cell = table.rows[r].cells[col_idx]
            text = cell.text.strip()
            logger.info("[_find_parent] up row=%d col=%d text=%r is_sub=%s",
                        r, col_idx, text[:20] if text else '', self._is_sub_item(text))
            if text and len(text) > 1 and not self._is_spec_text(text):
                if not self._is_sub_item(text):
                    logger.info("[_find_parent] found UP parent=%r", text)
                    return text

        logger.info("[_find_parent] no parent found for row=%d col=%d", row_idx, col_idx)
        return None

    def _get_header_row(self, table, num_rows: int, num_cols: int) -> tuple[int | None, int | None]:
        """找到数据表的表头行索引和检验项目列索引。

        检测报告通常结构：
        - 前几行是文档元数据（项目名称、委托单位等），列数不规则
        - 中间某行是数据表头（序、检测项目、单位、标准要求、检验结果、单项结论）
        - 后面是数据行

        返回 (header_row_idx, item_col_idx)。item_col_idx 为"检验项目"/"检测项目"列索引。
        """
        # 策略1：找含表头关键词的行（最可靠）
        HEADER_KEYWORDS = {"序号", "检验项目", "检测项目", "标准要求", "检验结果", "检测结果", "单项结论"}
        ITEM_KEYWORDS = {"检验项目", "检测项目"}
        for r in range(num_rows):
            row = table.rows[r]
            cell_texts = {}
            for c in range(num_cols):
                text = row.cells[c].text.strip().replace(" ", "")
                if text:
                    cell_texts[c] = text
            if set(cell_texts.values()) & HEADER_KEYWORDS:
                # 找到表头行，识别检验项目列
                item_col = None
                for c, text in cell_texts.items():
                    if text in ITEM_KEYWORDS:
                        item_col = c
                        break
                return r, item_col

        # 策略2：回退到启发式——所有列非空 + 文本短
        for r in range(num_rows):
            row = table.rows[r]
            non_empty = 0
            all_short = True
            has_long = False

            for c in range(num_cols):
                text = row.cells[c].text.strip()
                if text:
                    non_empty += 1
                    if len(text) > 6:
                        all_short = False
                    if len(text) > 20:
                        has_long = True

            if non_empty == num_cols and all_short and not has_long:
                return r, None

        return None, None

    def _is_spec_text(self, text: str) -> bool:
        """判断文本是否为规格要求、单位、日期等非检测项名称。"""
        import re
        # 规格要求：含比较符号（含全角 ± 等）
        if re.search(r'[≤≥<>=±＞﹤≧≦﹥]', text):
            return True
        # 纯数字+单位符号
        if re.match(r'^[\d.]+\s*[°%℃]', text):
            return True
        # 单位文本（如 W/(m·K)、MJ/kg、kg/m³、s、%），要求含特殊符号
        if re.match(r'^[A-Za-z°/%·\s()³²⁰-⁹]+$', text) and re.search(r'[°·³²℃]', text):
            return True
        if re.match(r'^[A-Za-z]+/[A-Za-z]', text) and len(text) <= 10:
            return True
        # 日期格式
        if re.match(r'^\d{4}[\.\-/]\d{1,2}[\.\-/]\d{1,2}', text):
            return True
        # 文档占位标记
        if text in ("以下空白", "以下为空白", "以下为空"):
            return True
        return False

    def _capture_format(self, cell) -> CellFormat:
        """捕获单元格的格式信息。"""
        fmt = CellFormat()

        if not cell.paragraphs:
            return fmt

        para = cell.paragraphs[0]
        fmt.alignment = para.alignment

        if para.runs:
            src = para.runs[0]
            fmt.bold = src.bold
            fmt.italic = src.italic
            fmt.underline = src.underline
            fmt.font_size = src.font.size
            fmt.font_name = src.font.name
            if src.font.color and src.font.color.rgb:
                fmt.color_rgb = src.font.color.rgb

        return fmt
