"""回填引擎：将提取结果按精确坐标写入 Word 模板，保留原始格式。"""

import os
import re
from docx import Document
from docx.oxml.ns import qn
from core.word_scanner import FieldSlot, CellFormat


class ReportFiller:
    """将数据按 FieldSlot 坐标精确回填到 Word 模板。"""

    def fill(self, template_path: str, output_path: str, slots: list[FieldSlot], data: dict[str, str | None]) -> list[str]:
        """回填数据到 Word 模板。

        Args:
            template_path: 模板文件路径
            output_path: 输出文件路径
            slots: WordScanner 扫描出的字段位置列表
            data: {标签名: 值} 字典

        Returns:
            填充日志列表
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"找不到模板文件: {template_path}")

        doc = Document(template_path)
        logs: list[str] = []

        for slot in slots:
            value = data.get(slot.label)
            if value is None:
                logs.append(f"跳过: [{slot.label}] — 未找到数据")
                continue

            try:
                table = doc.tables[slot.table_idx]
                row = table.rows[slot.row_idx]
                cell = row.cells[slot.cell_idx]
                self._write_to_cell(cell, str(value), slot.fmt)
                logs.append(f"填充: [{slot.label}] -> {value}")
            except (IndexError, KeyError) as e:
                logs.append(f"错误: [{slot.label}] — 坐标越界 {e}")

        doc.save(output_path)
        logs.append(f"已保存: {output_path}")
        return logs

    def _write_to_cell(self, cell, text: str, fmt: CellFormat):
        """写入数据并保留原有格式。"""
        if not cell.paragraphs:
            cell.text = text
            return

        para = cell.paragraphs[0]

        # 清除所有 run
        for run in para.runs:
            run._element.getparent().remove(run._element)

        # 创建新 run 并应用格式
        new_run = para.add_run(text)
        if fmt.bold is not None:
            new_run.bold = fmt.bold
        if fmt.italic is not None:
            new_run.italic = fmt.italic
        if fmt.underline is not None:
            new_run.underline = fmt.underline
        if fmt.font_size is not None:
            new_run.font.size = fmt.font_size
        if fmt.font_name is not None:
            new_run.font.name = fmt.font_name
            rpr = new_run._element.get_or_add_rPr()
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = rpr.makeelement(qn("w:rFonts"), {})
                rpr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), fmt.font_name)
        if fmt.color_rgb is not None:
            new_run.font.color.rgb = fmt.color_rgb

        # 恢复段落对齐
        if fmt.alignment is not None:
            para.alignment = fmt.alignment
